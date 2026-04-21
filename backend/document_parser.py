"""
文档解析模块 - 支持 Word、PDF、图片格式的需求文档解析
"""
import os
import io
import re
from typing import Dict, List, Optional, Tuple
from datetime import datetime

# 文档解析库
try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

try:
    import pypdf
except ImportError:
    pypdf = None

try:
    import pytesseract
    from PIL import Image
except ImportError:
    pytesseract = None
    Image = None


class DocumentParser:
    """文档解析器"""

    def __init__(self):
        self.supported_formats = ['.docx', '.pdf', '.png', '.jpg', '.jpeg', '.bmp', '.gif']

    def is_supported(self, file_path: str) -> bool:
        """检查文件格式是否支持"""
        ext = os.path.splitext(file_path)[1].lower()
        return ext in self.supported_formats

    def parse(self, file_path: str) -> Tuple[bool, str]:
        """
        解析文档

        Returns:
            (success, content_or_error)
        """
        if not os.path.exists(file_path):
            return False, f"文件不存在: {file_path}"

        ext = os.path.splitext(file_path)[1].lower()

        try:
            if ext == '.docx':
                return self._parse_docx(file_path)
            elif ext == '.pdf':
                return self._parse_pdf(file_path)
            elif ext in ['.png', '.jpg', '.jpeg', '.bmp', '.gif']:
                return self._parse_image(file_path)
            else:
                return False, f"不支持的文件格式: {ext}"
        except Exception as e:
            return False, f"解析失败: {str(e)}"

    def _parse_docx(self, file_path: str) -> Tuple[bool, str]:
        """解析 Word 文档"""
        if DocxDocument is None:
            return False, "未安装 python-docx 库，请运行: pip install python-docx"

        try:
            doc = DocxDocument(file_path)
            content_parts = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    content_parts.append(text)

            # 提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        content_parts.append(f"[表格] {row_text}")

            content = '\n'.join(content_parts)
            return True, content
        except Exception as e:
            return False, f"Word文档解析失败: {str(e)}"

    def _parse_pdf(self, file_path: str) -> Tuple[bool, str]:
        """解析 PDF 文档"""
        if pypdf is None:
            return False, "未安装 pypdf 库，请运行: pip install pypdf"

        try:
            reader = pypdf.PdfReader(file_path)
            content_parts = []

            for page_num, page in enumerate(reader.pages):
                text = page.extract_text()
                if text:
                    content_parts.append(f"[第{page_num + 1}页]\n{text}")

            content = '\n'.join(content_parts)
            return True, content
        except Exception as e:
            return False, f"PDF解析失败: {str(e)}"

    def _parse_image(self, file_path: str) -> Tuple[bool, str]:
        """解析图片 - 使用 OCR 提取文字"""
        if pytesseract is None or Image is None:
            return False, "未安装 OCR 库，请运行: pip install pytesseract pillow"

        try:
            image = Image.open(file_path)
            # 使用 OCR 提取文字
            text = pytesseract.image_to_string(image, lang='chi_sim+eng')

            # 清理文本
            text = self._clean_text(text)

            if not text.strip():
                return False, "图片中未识别到文字内容"

            return True, f"[图片内容]\n{text}"
        except Exception as e:
            return False, f"图片OCR解析失败: {str(e)}"

    def _clean_text(self, text: str) -> str:
        """清理OCR识别的文本"""
        # 移除多余空白
        text = re.sub(r'\s+', ' ', text)
        # 移除特殊字符
        text = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f]', '', text)
        return text.strip()


class RequirementAnalyzer:
    """需求分析器 - 从文档中提取结构化需求"""

    def __init__(self):
        self.parser = DocumentParser()

    def extract_requirements(self, content: str) -> List[Dict]:
        """
        从文档内容中提取需求

        Returns:
            需求列表，每个需求包含: title, description, acceptance_criteria
        """
        requirements = []

        # 按行分割内容
        lines = content.split('\n')

        current_req = None
        current_section = None
        current_content = []

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # 检测需求标题（通常以编号开头如 1.、1.1、需求一、REQ001 等）
            req_pattern = r'^(?:(\d+(?:\.\d+)*)[.、]\s*|REQ(\d+)[-:：]\s*|需求(\d+)[.、]\s*|第([一二三四五六七八九十]+)条\s*)(.+?)(?:\s*[（(].*?[）)])?$'

            # 检测验收标准（以"验收标准"、"通过条件"、"成功标准"开头）
            if any(kw in line for kw in ['验收标准', '通过条件', '成功标准', 'AC', 'AC:', 'AC-']):
                if current_req:
                    current_req['acceptance_criteria'] = line
                    continue

            # 检测功能描述开始
            if any(kw in line for kw in ['功能', '描述', '说明', '需求', '模块']):
                if current_req and current_content:
                    current_req['description'] = '\n'.join(current_content)
                    requirements.append(current_req)
                    current_content = []

                # 提取标题
                title = line
                match = re.match(req_pattern, line)
                if match:
                    title = match.group(5) if match.group(5) else line

                current_req = {
                    'title': title,
                    'description': '',
                    'acceptance_criteria': ''
                }
                continue

            if current_req is not None:
                current_content.append(line)

        # 处理最后一个需求
        if current_req and current_content:
            current_req['description'] = '\n'.join(current_content)
            requirements.append(current_req)

        # 如果没有识别到结构化需求，将整个内容作为一个需求
        if not requirements and content.strip():
            requirements.append({
                'title': '需求概述',
                'description': content.strip(),
                'acceptance_criteria': ''
            })

        return requirements


def parse_document(file_path: str) -> Tuple[bool, str]:
    """解析文档的便捷函数"""
    parser = DocumentParser()
    return parser.parse(file_path)


def extract_requirements_from_content(content: str) -> List[Dict]:
    """从文档内容中提取需求的便捷函数"""
    analyzer = RequirementAnalyzer()
    return analyzer.extract_requirements(content)
