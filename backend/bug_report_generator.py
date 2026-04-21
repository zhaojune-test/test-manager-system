"""
Bug报告生成器 - 根据导入的Bug数据生成测试报告Word文档
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os


class BugReportGenerator:
    """Bug测试报告生成器"""

    def __init__(self, template_path=None):
        self.template_path = template_path

    def set_cell_border(self, cell, **kwargs):
        """设置单元格边框"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for edge in ('top', 'left', 'bottom', 'right'):
            if edge in kwargs:
                element = OxmlElement(f'w:{edge}')
                element.set(qn('w:val'), 'single')
                element.set(qn('w:sz'), kwargs[edge])
                element.set(qn('w:space'), '0')
                element.set(qn('w:color'), '000000')
                tcBorders.append(element)
        tcPr.append(tcBorders)

    def generate_report(self, summary_data, output_path):
        """生成测试报告Word文档"""
        doc = Document()

        # 设置中文字体
        style = doc.styles['Normal']
        style.font.name = '宋体'
        style.font.size = Pt(12)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # 标题
        title = doc.add_paragraph()
        run = title.add_run(summary_data.get('project_name', 'EDSP-项目健康智能分析质量测试报告'))
        run.font.size = Pt(18)
        run.font.bold = True
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 开场白
        doc.add_paragraph('各位领导、同事好：')
        doc.add_paragraph(f'         {summary_data.get("project_name", "EDSP-项目健康智能分析")}已测试完成，现输出质量测试报告，请查收；如有问题及时沟通，谢谢。')

        # 测试结论
        conclusion_title = doc.add_paragraph()
        run = conclusion_title.add_run('测试结论：')
        run.font.bold = True

        conclusion = doc.add_paragraph()
        run = conclusion.add_run(summary_data.get('conclusion', '测试通过，已达交付标准。'))
        run.font.size = Pt(14)

        # 测试概况描述
        doc.add_paragraph()
        overview_title = doc.add_paragraph()
        run = overview_title.add_run('测试概况描述：')
        run.font.bold = True

        overview = doc.add_paragraph()
        run = overview.add_run(summary_data.get('overview', '本版本的整体提测质量优秀。本次需求均已按照需求评审时的需求实现，提测的需求内容开发思维逻辑完整，原型还原度较高，全流程无断点和阻塞。'))
        run.font.size = Pt(11)

        # 需要重点关注问题
        doc.add_paragraph()
        focus_title = doc.add_paragraph()
        run = focus_title.add_run('需要重点关注问题：')
        run.font.bold = True

        for item in summary_data.get('focus_points', []):
            doc.add_paragraph(f'         {item}')

        # 问题及风险
        doc.add_paragraph()
        risk_title = doc.add_paragraph()
        run = risk_title.add_run('问题及风险：')
        run.font.bold = True

        for item in summary_data.get('risks', []):
            doc.add_paragraph(f'         {item}')

        # 测试范围
        doc.add_paragraph()
        scope_title = doc.add_paragraph()
        run = scope_title.add_run('测试范围：')
        run.font.bold = True

        for item in summary_data.get('test_scope', []):
            doc.add_paragraph(f'         {item}')

        # Bug情况
        doc.add_paragraph()
        bug_title = doc.add_paragraph()
        run = bug_title.add_run('Bug情况：')
        run.font.bold = True

        doc.add_paragraph(f'截止目前，测试提出的有效bug总数{summary_data.get("total_bugs", 0)}个，均已回归验证通过。bug具体分类统计如下：')

        # 按bug类型统计表格
        doc.add_paragraph()
        category_title = doc.add_paragraph()
        run = category_title.add_run('按bug类型统计：')
        run.font.bold = True

        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '缺陷分类'
        hdr_cells[1].text = '数量'
        for category, count in summary_data.get('by_category', {}).items():
            row_cells = table.add_row().cells
            row_cells[0].text = str(category)
            row_cells[1].text = str(count)

        # 按bug严重程度统计表格
        doc.add_paragraph()
        severity_title = doc.add_paragraph()
        run = severity_title.add_run('按bug严重程度统计：')
        run.font.bold = True

        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '严重程度'
        hdr_cells[1].text = '数量'
        for severity, count in summary_data.get('by_severity', {}).items():
            row_cells = table.add_row().cells
            row_cells[0].text = str(severity)
            row_cells[1].text = str(count)

        # 按bug处理人统计表格
        doc.add_paragraph()
        handler_title = doc.add_paragraph()
        run = handler_title.add_run('按bug处理人统计：')
        run.font.bold = True

        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '处理人'
        hdr_cells[1].text = '数量'
        for handler, count in summary_data.get('by_handler', {}).items():
            row_cells = table.add_row().cells
            row_cells[0].text = str(handler)
            row_cells[1].text = str(count)

        # 按bug状态统计表格
        doc.add_paragraph()
        status_title = doc.add_paragraph()
        run = status_title.add_run('按bug状态统计：')
        run.font.bold = True

        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '状态'
        hdr_cells[1].text = '数量'
        for status, count in summary_data.get('by_status', {}).items():
            row_cells = table.add_row().cells
            row_cells[0].text = str(status)
            row_cells[1].text = str(count)

        # 测试时间
        doc.add_paragraph()
        time_title = doc.add_paragraph()
        run = time_title.add_run('测试时间：')
        run.font.bold = True

        test_time = doc.add_paragraph()
        run = test_time.add_run(summary_data.get('test_time', ''))

        # 产品、开发、测试人员
        doc.add_paragraph()
        team_title = doc.add_paragraph()
        run = team_title.add_run('产品、开发、测试人员：')
        run.font.bold = True

        for role, person in summary_data.get('team', {}).items():
            doc.add_paragraph(f'{role}：{person}')

        # 测试设备
        doc.add_paragraph()
        device_title = doc.add_paragraph()
        run = device_title.add_run('测试设备：')
        run.font.bold = True

        device = doc.add_paragraph()
        run = device.add_run(summary_data.get('test_device', '本轮测试所用浏览器：谷歌浏览器'))

        # 缺陷列表
        doc.add_paragraph()
        defect_list_title = doc.add_paragraph()
        run = defect_list_title.add_run('缺陷列表：')
        run.font.bold = True

        # 创建缺陷列表表格（添加原因列）
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '编号'
        hdr_cells[1].text = '缺陷标题'
        hdr_cells[2].text = '分类'
        hdr_cells[3].text = '严重程度'
        hdr_cells[4].text = '状态'
        hdr_cells[5].text = '缺陷原因'

        for defect in summary_data.get('defect_list', [])[:50]:  # 最多显示50条
            row_cells = table.add_row().cells
            row_cells[0].text = str(defect.get('id', ''))
            # 显示完整标题，不截断
            row_cells[1].text = str(defect.get('title', ''))
            row_cells[2].text = str(defect.get('category', ''))
            row_cells[3].text = str(defect.get('severity', ''))
            row_cells[4].text = str(defect.get('status', ''))
            # 显示缺陷原因
            reason = defect.get('reason') or defect.get('reason_category') or ''
            row_cells[5].text = reason

        # 附：质量维度评价标准
        doc.add_paragraph()
        doc.add_paragraph('附：质量维度评价标准')
        doc.add_paragraph('多个维度综合评价，取4个维度评价结果的众值作为最终评价结果，一票否决除外')

        # 保存文档
        doc.save(output_path)
        return output_path


def generate_test_report(bug_file_path, output_path, project_name='EDSP-项目健康智能分析', config=None):
    """生成测试报告的便捷函数"""
    from backend.defect_analyzer import DefectAnalyzer

    # 分析Bug数据
    analyzer = DefectAnalyzer()
    analyzer.load_excel(bug_file_path)
    summary = analyzer.get_summary()

    # 准备报告数据
    report_data = {
        'project_name': project_name,
        'conclusion': '测试通过，已达交付标准。',
        'overview': '本版本的整体提测质量优秀。本次需求均已按照需求评审时的需求实现，提测的需求内容开发思维逻辑完整，原型还原度较高，全流程无断点和阻塞。但是由于涉及的功能点较多较细，会存在一些功能细节的忽略，导致出现一些bug。在测试中也存在需求评审时未考虑到的个别细节问题。',
        'focus_points': [
            '项目健康智能分析每天定时任务更新数据；',
            '上线前需提前同步线上数据库表结构变化、定时任务配置和权限配置等'
        ],
        'risks': [
            '由于项目健康智能分析数据是每天定时任务更新数据，可能会存在用户及时处理了项目如编制了项目计划、预算等，却看到分数和风险等数据未实时更新，需要和用户表明，是第二天定时任务更新数据，不是实时更新。',
            '本次需求功能均是在测试环境中测试实现，测试环境的数据是认为构造，待生产发布后，可能会有某个细节点的不一致，需要及时关注线上相关功能情况。'
        ],
        'test_scope': [
            '项目详情页的项目总结展示、项目得分展示；',
            '项目智能分析报告的雷达图显示，五个维度：项目损益、客户满意、项目进度、项目风险/质量和供应商管理的数据展示；',
            '项目损益、客户满意、项目进度、项目风险/质量和供应商管理的得分逻辑、自动罗列问题展示、自动创建风险、自动创建问题、及自动闭环风险和问题等；',
            '完整版项目分析报告的数据展示和导出功能；',
            'PMO及内控、用户体验李鹏，采购于敬凤等人员的项目新增亮点功能、手动创建问题同步健康报告的功能；',
            '订单项目和生产交付类项目的得分逻辑计算'
        ],
        'total_bugs': summary['total'],
        'by_category': summary['by_category'],
        'by_severity': summary['by_severity'],
        'by_handler': summary['by_handler'],
        'by_status': summary['by_status'],
        'test_time': '2025.02.28—2025.03.05',
        'team': {
            '产品经理': '余艳红',
            '后端': '徐见欢、郑云（项目损益中涉及四算功能的为郑云同学）',
            '前端': '王彦超',
            '测试': '赵俊娥'
        },
        'test_device': '本轮测试所用浏览器：谷歌浏览器',
        'defect_list': analyzer.get_defect_list()
    }

    # 如果有自定义配置，覆盖默认数据
    if config:
        report_data.update(config)

    # 生成报告
    generator = BugReportGenerator()
    generator.generate_report(report_data, output_path)

    return report_data


if __name__ == '__main__':
    # 测试生成报告
    result = generate_test_report(
        r'E:\June_workspace\oms\test_manager\缺陷信息(1).xlsx',
        r'E:\June_workspace\oms\test_manager\测试报告_output.docx'
    )
    print(f"报告生成成功: {result['total_bugs']} 个缺陷")
