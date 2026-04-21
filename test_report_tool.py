"""
测试报告生成工具 - 独立桌面应用程序
无需Flask服务器，直接运行即可使用
"""
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import os
import re
import pandas as pd
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class TestReportTool:
    """测试报告生成工具"""

    def __init__(self, root):
        self.root = root
        self.root.title("测试报告生成工具")
        self.root.geometry("900x700")

        # 数据变量
        self.excel_path = tk.StringVar()
        self.df = None
        self.summary_data = {}

        # 创建界面
        self.create_widgets()

    def create_widgets(self):
        """创建界面组件"""
        # 主框架
        main_frame = ttk.Frame(self.root, padding="10")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))

        # 文件选择区域
        file_frame = ttk.LabelFrame(main_frame, text="1. 选择缺陷Excel文件", padding="10")
        file_frame.grid(row=0, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=5)

        ttk.Entry(file_frame, textvariable=self.excel_path, width=60).grid(row=0, column=0, padx=5)
        ttk.Button(file_frame, text="浏览...", command=self.browse_file).grid(row=0, column=1)
        ttk.Button(file_frame, text="导入数据", command=self.import_data).grid(row=0, column=2, padx=5)

        # 统计信息显示
        stats_frame = ttk.LabelFrame(main_frame, text="2. 缺陷统计信息", padding="10")
        stats_frame.grid(row=1, column=0, columnspan=2, sticky=(tk.W, tk.E, tk.N, tk.S), pady=5)

        self.stats_text = tk.Text(stats_frame, height=8, width=80)
        self.stats_text.grid(row=0, column=0, sticky=(tk.W, tk.E))

        scrollbar = ttk.Scrollbar(stats_frame, orient=tk.VERTICAL, command=self.stats_text.yview)
        scrollbar.grid(row=0, column=1, sticky=(tk.N, tk.S))
        self.stats_text.config(yscrollcommand=scrollbar.set)

        # 配置区域
        config_frame = ttk.LabelFrame(main_frame, text="3. 报告配置", padding="10")
        config_frame.grid(row=2, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=5)

        # 项目名称
        ttk.Label(config_frame, text="项目名称:").grid(row=0, column=0, sticky=tk.W, pady=3)
        self.project_name = ttk.Entry(config_frame, width=50)
        self.project_name.grid(row=0, column=1, sticky=(tk.W, tk.E), pady=3)

        # 测试时间
        ttk.Label(config_frame, text="测试时间:").grid(row=1, column=0, sticky=tk.W, pady=3)
        self.test_time = ttk.Entry(config_frame, width=50)
        self.test_time.grid(row=1, column=1, sticky=(tk.W, tk.E), pady=3)

        # 测试结论
        ttk.Label(config_frame, text="测试结论:").grid(row=2, column=0, sticky=tk.W, pady=3)
        self.conclusion = ttk.Entry(config_frame, width=50)
        self.conclusion.grid(row=2, column=1, sticky=(tk.W, tk.E), pady=3)

        # 测试概况描述
        ttk.Label(config_frame, text="测试概况描述:").grid(row=3, column=0, sticky=tk.W, pady=3)
        self.overview = tk.Text(config_frame, height=3, width=38)
        self.overview.grid(row=3, column=1, sticky=(tk.W, tk.E), pady=3)

        # 需要重点关注问题
        ttk.Label(config_frame, text="重点关注问题:").grid(row=4, column=0, sticky=tk.W, pady=3)
        self.focus_points = tk.Text(config_frame, height=3, width=38)
        self.focus_points.grid(row=4, column=1, sticky=(tk.W, tk.E), pady=3)

        # 问题及风险
        ttk.Label(config_frame, text="问题及风险:").grid(row=5, column=0, sticky=tk.W, pady=3)
        self.risks = tk.Text(config_frame, height=3, width=38)
        self.risks.grid(row=5, column=1, sticky=(tk.W, tk.E), pady=3)

        # 测试范围
        ttk.Label(config_frame, text="测试范围:").grid(row=6, column=0, sticky=tk.W, pady=3)
        self.test_scope = tk.Text(config_frame, height=3, width=38)
        self.test_scope.grid(row=6, column=1, sticky=(tk.W, tk.E), pady=3)

        # 团队配置
        team_frame = ttk.LabelFrame(main_frame, text="4. 团队成员配置", padding="10")
        team_frame.grid(row=3, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=5)

        ttk.Label(team_frame, text="产品经理:").grid(row=0, column=0, sticky=tk.W, pady=3)
        self.team_product = ttk.Entry(team_frame, width=30)
        self.team_product.grid(row=0, column=1, sticky=(tk.W, tk.E), pady=3, padx=5)

        ttk.Label(team_frame, text="后端开发:").grid(row=0, column=2, sticky=tk.W, pady=3)
        self.team_backend = ttk.Entry(team_frame, width=30)
        self.team_backend.grid(row=0, column=3, sticky=(tk.W, tk.E), pady=3, padx=5)

        ttk.Label(team_frame, text="前端开发:").grid(row=1, column=0, sticky=tk.W, pady=3)
        self.team_frontend = ttk.Entry(team_frame, width=30)
        self.team_frontend.grid(row=1, column=1, sticky=(tk.W, tk.E), pady=3, padx=5)

        ttk.Label(team_frame, text="测试人员:").grid(row=1, column=2, sticky=tk.W, pady=3)
        self.team_tester = ttk.Entry(team_frame, width=30)
        self.team_tester.grid(row=1, column=3, sticky=(tk.W, tk.E), pady=3, padx=5)

        # 测试设备
        ttk.Label(team_frame, text="测试设备:").grid(row=2, column=0, sticky=tk.W, pady=3)
        self.test_device = ttk.Entry(team_frame, width=30)
        self.test_device.grid(row=2, column=1, sticky=(tk.W, tk.E), pady=3, padx=5)

        # 生成按钮
        btn_frame = ttk.Frame(main_frame)
        btn_frame.grid(row=4, column=0, columnspan=2, pady=10)

        ttk.Button(btn_frame, text="自动填充配置", command=self.auto_fill_config).grid(row=0, column=0, padx=5)
        ttk.Button(btn_frame, text="清空配置", command=self.clear_config).grid(row=0, column=1, padx=5)
        ttk.Button(btn_frame, text="生成报告", command=self.generate_report, style="Accent.TButton").grid(row=0, column=2, padx=5)

        # 进度显示
        self.progress = ttk.Progressbar(main_frame, mode='indeterminate')
        self.progress.grid(row=5, column=0, columnspan=2, sticky=(tk.W, tk.E))

    def browse_file(self):
        """浏览选择Excel文件"""
        file_path = filedialog.askopenfilename(
            title="选择缺陷Excel文件",
            filetypes=[("Excel文件", "*.xlsx *.xls"), ("所有文件", "*.*")]
        )
        if file_path:
            self.excel_path.set(file_path)

    def import_data(self):
        """导入Excel数据"""
        file_path = self.excel_path.get()
        if not file_path:
            messagebox.showwarning("警告", "请先选择Excel文件")
            return

        if not os.path.exists(file_path):
            messagebox.showerror("错误", "文件不存在")
            return

        try:
            self.df = pd.read_excel(file_path)
            self.analyze_data()
            messagebox.showinfo("成功", f"成功导入 {len(self.df)} 条缺陷数据")
        except Exception as e:
            messagebox.showerror("错误", f"导入失败: {str(e)}")

    def analyze_data(self):
        """分析数据并显示统计信息"""
        if self.df is None:
            return

        total = len(self.df)
        stats = []

        # 按缺陷分类统计
        if '缺陷分类' in self.df.columns:
            by_category = self.df['缺陷分类'].dropna().value_counts()
            stats.append(f"缺陷分类: {dict(by_category)}")

        # 按严重程度统计
        if '严重程度' in self.df.columns:
            by_severity = self.df['严重程度'].dropna().value_counts()
            stats.append(f"严重程度: {dict(by_severity)}")

        # 按处理人统计
        if '处理人' in self.df.columns:
            by_handler = self.df['处理人'].dropna().value_counts()
            stats.append(f"处理人: {dict(by_handler)}")

        # 按状态统计
        if '状态' in self.df.columns:
            by_status = self.df['状态'].dropna().value_counts()
            stats.append(f"状态: {dict(by_status)}")

        # 获取测试时间范围
        test_time_str = ""
        if '关闭时间' in self.df.columns:
            dates = self.df['关闭时间'].dropna()
            if len(dates) > 0:
                test_start = pd.to_datetime(dates.min())
                test_end = pd.to_datetime(dates.max())
                test_time_str = f"{test_start.strftime('%Y.%m.%d')}—{test_end.strftime('%Y.%m.%d')}"
                stats.append(f"测试时间: {test_time_str}")

        # 获取版本号
        if '版本' in self.df.columns:
            versions = self.df['版本'].dropna().unique()
            if len(versions) > 0:
                stats.append(f"版本: {versions[0]}")

        # 获取验证人
        verifiers = []
        if '验证人' in self.df.columns:
            verifiers = self.df['验证人'].dropna().unique().tolist()
            stats.append(f"验证人: {verifiers}")

        # 获取测试范围（从标题中提取【】内容）
        scopes = set()
        if '标题' in self.df.columns:
            for title in self.df['标题'].dropna():
                contents = re.findall(r'【([^】]+)】', str(title))
                for content in contents:
                    if content and len(content) < 20:
                        scopes.add(content)
        if scopes:
            stats.append(f"测试范围: {list(scopes)}")

        # 显示统计信息
        self.stats_text.delete(1.0, tk.END)
        self.stats_text.insert(tk.END, f"总缺陷数: {total}\n\n")
        for s in stats:
            self.stats_text.insert(tk.END, f"{s}\n\n")

        # 保存摘要数据供自动填充使用
        self.summary_data = {
            'total': total,
            'by_category': by_category.to_dict() if '缺陷分类' in self.df.columns else {},
            'by_severity': by_severity.to_dict() if '严重程度' in self.df.columns else {},
            'by_handler': by_handler.to_dict() if '处理人' in self.df.columns else {},
            'by_status': by_status.to_dict() if '状态' in self.df.columns else {},
            'test_time': test_time_str,
            'version': versions[0] if '版本' in self.df.columns and len(versions) > 0 else '',
            'verifiers': verifiers,
            'scopes': list(scopes)
        }

    def auto_fill_config(self):
        """自动填充配置"""
        if not self.summary_data:
            messagebox.showwarning("警告", "请先导入Excel数据")
            return

        # 自动填充项目名称（使用版本号）
        if self.summary_data.get('version'):
            self.project_name.set(f"{self.summary_data['version']}测试报告")

        # 自动填充测试时间
        if self.summary_data.get('test_time'):
            self.test_time.set(self.summary_data['test_time'])

        # 自动填充测试结论
        self.conclusion.set("测试通过，已达交付标准。")

        # 自动填充测试概况描述
        overview = "本版本的整体提测质量优秀。本次需求均已按照需求评审时的需求实现，提测的需求内容开发思维逻辑完整，原型还原度较高，全流程无断点和阻塞。"
        self.overview.delete(1.0, tk.END)
        self.overview.insert(tk.END, overview)

        # 自动填充重点关注问题
        focus = "1. 定时任务更新数据需关注；\n2. 上线前需提前同步线上配置"
        self.focus_points.delete(1.0, tk.END)
        self.focus_points.insert(tk.END, focus)

        # 自动填充问题及风险
        risk = "1. 数据更新为定时任务，非实时更新；\n2. 测试环境数据为构造数据"
        self.risks.delete(1.0, tk.END)
        self.risks.insert(tk.END, risk)

        # 自动填充测试范围
        scopes = self.summary_data.get('scopes', [])
        if scopes:
            scope_text = "；".join(scopes)
            self.test_scope.delete(1.0, tk.END)
            self.test_scope.insert(tk.END, scope_text)

        # 自动填充测试人员
        verifiers = self.summary_data.get('verifiers', [])
        if verifiers:
            self.team_tester.set("、".join([str(v) for v in verifiers[:3]]))

        messagebox.showinfo("成功", "配置已自动填充")

    def clear_config(self):
        """清空配置"""
        self.project_name.set("")
        self.test_time.set("")
        self.conclusion.set("")
        self.overview.delete(1.0, tk.END)
        self.focus_points.delete(1.0, tk.END)
        self.risks.delete(1.0, tk.END)
        self.test_scope.delete(1.0, tk.END)
        self.team_product.set("")
        self.team_backend.set("")
        self.team_frontend.set("")
        self.team_tester.set("")
        self.test_device.set("")

    def generate_report(self):
        """生成测试报告"""
        if self.df is None:
            messagebox.showwarning("警告", "请先导入Excel数据")
            return

        # 获取保存路径
        default_name = f"测试报告_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        save_path = filedialog.asksaveasfilename(
            title="保存测试报告",
            defaultextension=".docx",
            filetypes=[("Word文档", "*.docx"), ("所有文件", "*.*")],
            initialfile=default_name
        )

        if not save_path:
            return

        self.progress.start()

        try:
            # 获取配置数据
            report_data = {
                'project_name': self.project_name.get() or '测试报告',
                'conclusion': self.conclusion.get() or '测试通过',
                'overview': self.overview.get(1.0, tk.END).strip(),
                'focus_points': self.focus_points.get(1.0, tk.END).strip().split('\n'),
                'risks': self.risks.get(1.0, tk.END).strip().split('\n'),
                'test_scope': self.test_scope.get(1.0, tk.END).strip().split('\n'),
                'total_bugs': self.summary_data.get('total', len(self.df)),
                'by_category': self.summary_data.get('by_category', {}),
                'by_severity': self.summary_data.get('by_severity', {}),
                'by_handler': self.summary_data.get('by_handler', {}),
                'by_status': self.summary_data.get('by_status', {}),
                'test_time': self.test_time.get() or '',
                'team': {
                    '产品经理': self.team_product.get(),
                    '后端': self.team_backend.get(),
                    '前端': self.team_frontend.get(),
                    '测试': self.team_tester.get()
                },
                'test_device': self.test_device.get() or '谷歌浏览器',
                'defect_list': self.get_defect_list()
            }

            # 清理空数据
            report_data['focus_points'] = [p.strip() for p in report_data['focus_points'] if p.strip()]
            report_data['risks'] = [r.strip() for r in report_data['risks'] if r.strip()]
            report_data['test_scope'] = [s.strip() for s in report_data['test_scope'] if s.strip()]

            # 生成Word文档
            self.create_word_report(report_data, save_path)

            self.progress.stop()
            messagebox.showinfo("成功", f"报告已生成: {save_path}")

            # 清空配置
            self.clear_config()

        except Exception as e:
            self.progress.stop()
            messagebox.showerror("错误", f"生成失败: {str(e)}")

    def get_defect_list(self):
        """获取缺陷列表"""
        if self.df is None:
            return []

        result = []
        for _, row in self.df.iterrows():
            def safe_get(key, default=''):
                val = row.get(key, default)
                if pd.notna(val):
                    return str(val)
                return ''

            result.append({
                'id': safe_get('编号'),
                'title': safe_get('标题'),
                'category': safe_get('缺陷分类'),
                'severity': safe_get('严重程度'),
                'status': safe_get('状态'),
                'reason': safe_get('缺陷具体原因'),
                'reason_category': safe_get('缺陷原因分类')
            })
        return result

    def create_word_report(self, data, output_path):
        """创建Word报告"""
        doc = Document()

        # 设置中文字体
        style = doc.styles['Normal']
        style.font.name = '宋体'
        style.font.size = Pt(12)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # 标题
        title = doc.add_paragraph()
        run = title.add_run(data.get('project_name', '测试报告'))
        run.font.size = Pt(18)
        run.font.bold = True
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 开场白
        doc.add_paragraph('各位领导、同事好：')
        doc.add_paragraph(f"         {data.get('project_name', '项目')}已测试完成，现输出质量测试报告，请查收；如有问题及时沟通，谢谢。")

        # 测试结论
        conclusion_title = doc.add_paragraph()
        run = conclusion_title.add_run('测试结论：')
        run.font.bold = True

        conclusion = doc.add_paragraph()
        run = conclusion.add_run(data.get('conclusion', '测试通过'))
        run.font.size = Pt(14)

        # 测试概况描述
        doc.add_paragraph()
        overview_title = doc.add_paragraph()
        run = overview_title.add_run('测试概况描述：')
        run.font.bold = True

        overview = doc.add_paragraph()
        run = overview.add_run(data.get('overview', ''))
        run.font.size = Pt(11)

        # 需要重点关注问题
        doc.add_paragraph()
        focus_title = doc.add_paragraph()
        run = focus_title.add_run('需要重点关注问题：')
        run.font.bold = True

        for item in data.get('focus_points', []):
            doc.add_paragraph(f'         {item}')

        # 问题及风险
        doc.add_paragraph()
        risk_title = doc.add_paragraph()
        run = risk_title.add_run('问题及风险：')
        run.font.bold = True

        for item in data.get('risks', []):
            doc.add_paragraph(f'         {item}')

        # 测试范围
        doc.add_paragraph()
        scope_title = doc.add_paragraph()
        run = scope_title.add_run('测试范围：')
        run.font.bold = True

        for item in data.get('test_scope', []):
            doc.add_paragraph(f'         {item}')

        # Bug情况
        doc.add_paragraph()
        bug_title = doc.add_paragraph()
        run = bug_title.add_run('Bug情况：')
        run.font.bold = True

        doc.add_paragraph(f"截止目前，测试提出的有效bug总数{data.get('total_bugs', 0)}个，均已回归验证通过。bug具体分类统计如下：")

        # 按bug类型统计表格
        doc.add_paragraph()
        category_title = doc.add_paragraph()
        run = category_title.add_run('按bug类型统计：')
        run.font.bold = True

        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '缺陷分类'
        hdr_cells[1].text = '数量'
        for category, count in data.get('by_category', {}).items():
            row_cells = table.add_row().cells
            row_cells[0].text = str(category)
            row_cells[1].text = str(count)

        # 按bug严重程度统计表格
        doc.add_paragraph()
        severity_title = doc.add_paragraph()
        run = severity_title.add_run('按bug严重程度统计：')
        run.font.bold = True

        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '严重程度'
        hdr_cells[1].text = '数量'
        for severity, count in data.get('by_severity', {}).items():
            row_cells = table.add_row().cells
            row_cells[0].text = str(severity)
            row_cells[1].text = str(count)

        # 按bug处理人统计表格
        doc.add_paragraph()
        handler_title = doc.add_paragraph()
        run = handler_title.add_run('按bug处理人统计：')
        run.font.bold = True

        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '处理人'
        hdr_cells[1].text = '数量'
        for handler, count in data.get('by_handler', {}).items():
            row_cells = table.add_row().cells
            row_cells[0].text = str(handler)
            row_cells[1].text = str(count)

        # 按bug状态统计表格
        doc.add_paragraph()
        status_title = doc.add_paragraph()
        run = status_title.add_run('按bug状态统计：')
        run.font.bold = True

        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '状态'
        hdr_cells[1].text = '数量'
        for status, count in data.get('by_status', {}).items():
            row_cells = table.add_row().cells
            row_cells[0].text = str(status)
            row_cells[1].text = str(count)

        # 测试时间
        doc.add_paragraph()
        time_title = doc.add_paragraph()
        run = time_title.add_run('测试时间：')
        run.font.bold = True

        test_time = doc.add_paragraph()
        run = test_time.add_run(data.get('test_time', ''))

        # 产品、开发、测试人员
        doc.add_paragraph()
        team_title = doc.add_paragraph()
        run = team_title.add_run('产品、开发、测试人员：')
        run.font.bold = True

        team = data.get('team', {})
        for role, person in team.items():
            if person:
                doc.add_paragraph(f'{role}：{person}')

        # 测试设备
        doc.add_paragraph()
        device_title = doc.add_paragraph()
        run = device_title.add_run('测试设备：')
        run.font.bold = True

        device = doc.add_paragraph()
        run = device.add_run(f"本轮测试所用浏览器：{data.get('test_device', '谷歌浏览器')}")

        # 缺陷列表
        doc.add_paragraph()
        defect_list_title = doc.add_paragraph()
        run = defect_list_title.add_run('缺陷列表：')
        run.font.bold = True

        # 创建缺陷列表表格
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '编号'
        hdr_cells[1].text = '缺陷标题'
        hdr_cells[2].text = '分类'
        hdr_cells[3].text = '严重程度'
        hdr_cells[4].text = '状态'
        hdr_cells[5].text = '缺陷原因'

        for defect in data.get('defect_list', [])[:50]:
            row_cells = table.add_row().cells
            row_cells[0].text = str(defect.get('id', ''))
            row_cells[1].text = str(defect.get('title', ''))
            row_cells[2].text = str(defect.get('category', ''))
            row_cells[3].text = str(defect.get('severity', ''))
            row_cells[4].text = str(defect.get('status', ''))
            reason = defect.get('reason') or defect.get('reason_category') or ''
            row_cells[5].text = reason

        # 保存文档
        doc.save(output_path)


def main():
    """主函数"""
    root = tk.Tk()
    app = TestReportTool(root)
    root.mainloop()


if __name__ == '__main__':
    main()
